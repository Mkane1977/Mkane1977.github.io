import sys
import re
import spacy
import string
from collections import Counter
from PyQt5.QtGui import QTextCharFormat, QColor
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QTextEdit,
    QPushButton, QCheckBox, QHBoxLayout, QLabel, QComboBox, QFileDialog
)
from docx import Document
import PyPDF2
from spacy.tokens import Span

# Matplotlib for charts (Phase III stats visualization)
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

#############################################################
# 1) Define the set of standard spaCy entity labels to display #
#############################################################
CHECKBOX_ENTITY_LABELS = [
    "PERSON",
    "ORG",
    "GPE",
    "DATE",
    "TIME",
    "PRODUCT",
    "MONEY",
    "QUANTITY",
    "PHONE",
    "EMAIL",
    "URL",
    "PERCENT",
]

##########################################
# 2) Map each entity label to a color        #
# Adjust these hex values to match your UI #
##########################################
ENTITY_COLORS = {
    "PERSON": "#C8E6C9",    # Light green
    "ORG": "#A2C8EC",       # Light blue
    "GPE": "#FFCDD2",       # Light red
    "DATE": "#FFF9C4",      # Light yellow
    "TIME": "#F8BBD0",      # Light pink
    "PRODUCT": "#B2EBF2",   # Light cyan
    "MONEY": "#D1C4E9",     # Light purple
    "QUANTITY": "#C5CAE9",  # Indigo variant
    "PHONE": "#E0F7FA",     # Very light cyan
    "EMAIL": "#E1BEE7",     # Lilac
    "URL": "#DCEDC8",       # Light green variant
    "PERCENT": "#FFCCBC"    # Light orange
}

def preprocess_text(text):
    """
    Insert periods in text without punctuation to help spaCy with sentence segmentation.
    Specifically, after a recognized date/price if the next token starts with uppercase.
    """
    tokens = text.split()
    new_tokens = []
    for i in range(len(tokens)):
        new_tokens.append(tokens[i])
        if i < len(tokens) - 1:
            if (re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}$", tokens[i]) or
                re.match(r"^\$?\d+(\.\d+)?$", tokens[i])):
                if tokens[i+1] and tokens[i+1][0].isupper():
                    new_tokens.append(".")
    return " ".join(new_tokens)

def set_run_background(run, color_hex):
    """
    Apply a background color (shading) to a run in the Word document using its underlying XML.
    Removes a leading '#' if present.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), "clear")
    color_hex = color_hex.lstrip('#')
    shd.set(qn('w:fill'), color_hex)
    run._element.get_or_add_rPr().append(shd)

class NERApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Named Entity Recognizer - Phase III")
        self.spacy_models = ["en_core_web_sm", "en_core_web_md", "en_core_web_lg"]
        self.selected_entities = []
        try:
            self.nlp = spacy.load(self.spacy_models[0])
        except OSError:
            raise SystemExit(
                "SpaCy model not found. Please run:\n\n   python -m spacy download en_core_web_sm\n"
            )
        self.current_model = self.spacy_models[0]
        self.initUI()

    def initUI(self):
        """Construct the complete PyQt interface."""
        layout = QVBoxLayout()

        # Model selection dropdown
        self.model_dropdown = QComboBox(self)
        self.model_dropdown.addItems(self.spacy_models)
        self.model_dropdown.currentTextChanged.connect(self.update_model)
        layout.addWidget(QLabel("Select spaCy Model:"))
        layout.addWidget(self.model_dropdown)

        # Input text area for raw text or loaded file content
        self.text_input = QTextEdit(self)
        layout.addWidget(QLabel("Input Text (or Loaded File):"))
        layout.addWidget(self.text_input)

        # Button to load a file (Word or PDF)
        self.open_btn = QPushButton("Open File", self)
        self.open_btn.clicked.connect(self.load_file)
        layout.addWidget(self.open_btn)

        # Label for the entity checkboxes section
        layout.addWidget(QLabel("Entity Labels:"))

        # Buttons for "Select All" and "Deselect All"
        button_layout = QHBoxLayout()
        select_all_btn = QPushButton("Select All", self)
        select_all_btn.clicked.connect(self.select_all_entities)
        button_layout.addWidget(select_all_btn)
        deselect_all_btn = QPushButton("Deselect All", self)
        deselect_all_btn.clicked.connect(self.deselect_all_entities)
        button_layout.addWidget(deselect_all_btn)
        layout.addLayout(button_layout)

        # Container for entity checkboxes (horizontal layout)
        self.checkboxes = {}
        checkbox_layout = QHBoxLayout()
        for label in CHECKBOX_ENTITY_LABELS:
            cb = QCheckBox(label, self)
            # Initially, unchecked: background is transparent
            cb.setStyleSheet("QCheckBox { background-color: transparent; padding: 3px; }")
            cb.stateChanged.connect(self.update_selected_entities)
            self.checkboxes[label] = cb
            checkbox_layout.addWidget(cb)
        layout.addLayout(checkbox_layout)

        # Button to run NER tagging
        self.tag_btn = QPushButton("Tag Entities", self)
        self.tag_btn.clicked.connect(self.tag_entities)
        layout.addWidget(self.tag_btn)

        # Buttons to save output: Word and PDF
        save_button_layout = QHBoxLayout()
        self.save_word_btn = QPushButton("Save as Word", self)
        self.save_word_btn.clicked.connect(self.save_to_word)
        save_button_layout.addWidget(self.save_word_btn)
        self.save_pdf_btn = QPushButton("Save as PDF", self)
        self.save_pdf_btn.clicked.connect(self.save_to_pdf)
        save_button_layout.addWidget(self.save_pdf_btn)
        layout.addLayout(save_button_layout)

        # Output text area for the tagged text
        self.text_output = QTextEdit(self)
        self.text_output.setReadOnly(True)
        layout.addWidget(QLabel("Tagged Output:"))
        layout.addWidget(self.text_output)

        # --- PHASE III Additions ---
        # STAT button to generate document statistics
        self.stat_btn = QPushButton("STAT", self)
        self.stat_btn.clicked.connect(self.show_statistics)
        layout.addWidget(self.stat_btn)

        # Container for displaying statistics (numeric text + charts)
        self.stats_container = QWidget(self)
        self.stats_layout = QVBoxLayout(self.stats_container)
        layout.addWidget(self.stats_container)
        # ---------------------------

        # Set the main layout of the application window
        central_widget = QWidget(self)
        central_widget.setLayout(layout)
        self.setCentralWidget(central_widget)

    def select_all_entities(self):
        """Check all entity checkboxes and update selection and style."""
        for cb in self.checkboxes.values():
            cb.setChecked(True)
        self.update_selected_entities()

    def deselect_all_entities(self):
        """Uncheck all entity checkboxes and update selection and style."""
        for cb in self.checkboxes.values():
            cb.setChecked(False)
        self.update_selected_entities()

    def update_selected_entities(self):
        """Update the list of selected entities and update checkbox styling."""
        self.selected_entities = [
            label for label, cb in self.checkboxes.items() if cb.isChecked()
        ]
        self.update_checkbox_styles()

    def update_checkbox_styles(self):
        """
        Update each checkboxes style:
          - When checked, its background is the entity color.
          - When unchecked, its background is transparent.
        """
        for label, cb in self.checkboxes.items():
            if cb.isChecked():
                cb.setStyleSheet(
                    f"QCheckBox {{ background-color: {ENTITY_COLORS.get(label, '#FFFF00')}; padding: 3px; border-radius: 5px; color: black; }}"
                )
            else:
                cb.setStyleSheet("QCheckBox { background-color: transparent; padding: 3px; }")

    def update_model(self, model_name):
        """Switch to the selected spaCy model."""
        try:
            self.nlp = spacy.load(model_name)
            self.current_model = model_name
        except OSError as e:
            print(f"Could not load spaCy model '{model_name}': {e}")

    def load_file(self):
        """Load text from a Word or PDF file into the input text area."""
        file_dialog_opts = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Word Files (*.docx);;PDF Files (*.pdf)",
            options=file_dialog_opts
        )
        if not file_path:
            return
        loaded_text = ""
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            loaded_text = "\n".join(p.text for p in doc.paragraphs)
        elif file_path.endswith(".pdf"):
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                loaded_text = "\n".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
        self.text_input.setPlainText(loaded_text)

    def tag_entities(self):
        """
        Perform NER tagging on the input text.
        Only entities whose label is selected (checkbox checked) are highlighted.
        """
        text = self.text_input.toPlainText().strip()
        if not text:
            return
        if not any(ch in text for ch in string.punctuation):
            text = preprocess_text(text)
        doc_nlp = self.nlp(text)
        filtered_ents = [ent for ent in doc_nlp.ents if ent.label_ in self.selected_entities]
        doc_nlp.ents = filtered_ents
        self.highlight_entities(text, doc_nlp.ents)

    def highlight_entities(self, text, entities):
        """Display the tagged output with color highlights for entities."""
        self.text_output.clear()
        cursor = self.text_output.textCursor()
        last_idx = 0
        for ent in entities:
            cursor.insertText(text[last_idx:ent.start_char])
            fmt = QTextCharFormat()
            color = ENTITY_COLORS.get(ent.label_, "#FFFF00")
            fmt.setBackground(QColor(color))
            cursor.insertText(f"{ent.text} [{ent.label_}]", fmt)
            last_idx = ent.end_char
        cursor.insertText(text[last_idx:])

    def save_to_word(self):
        """
        Save the tagged output to a Word document.
        This version preserves entity background color by splitting text
        into runs and using custom XML shading.
        """
        file_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Files (*.docx)")
        if not file_path:
            return
        text = self.text_input.toPlainText().strip()
        if not text:
            return
        if not any(ch in text for ch in string.punctuation):
            text = preprocess_text(text)
        doc_nlp = self.nlp(text)
        filtered_ents = [ent for ent in doc_nlp.ents if ent.label_ in self.selected_entities]
        doc_nlp.ents = filtered_ents
        new_doc = Document()
        paragraph = new_doc.add_paragraph()
        last_idx = 0
        for ent in doc_nlp.ents:
            plain_segment = text[last_idx:ent.start_char]
            if plain_segment:
                paragraph.add_run(plain_segment)
            ent_segment = text[ent.start_char:ent.end_char]
            if ent_segment:
                run = paragraph.add_run(ent_segment + f" [{ent.label_}]")
                set_run_background(run, ENTITY_COLORS.get(ent.label_, "FFFF00"))
            last_idx = ent.end_char
        if last_idx < len(text):
            paragraph.add_run(text[last_idx:])
        new_doc.save(file_path)

    def save_to_pdf(self):
        """
        Save the tagged output to a PDF file with entity background highlighting.
        This is done using ReportLab's Platypus and inline markup via the <span> tag.
        """
        file_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF Files (*.pdf)")
        if not file_path:
            return
        text = self.text_input.toPlainText().strip()
        if not text:
            return
        if not any(ch in text for ch in string.punctuation):
            text = preprocess_text(text)
        doc_nlp = self.nlp(text)
        filtered_ents = [ent for ent in doc_nlp.ents if ent.label_ in self.selected_entities]
        doc_nlp.ents = filtered_ents
        # Build an HTML-like string with inline <span> tags using ReportLab markup:
        pdf_text = ""
        last_idx = 0
        for ent in doc_nlp.ents:
            plain_segment = text[last_idx:ent.start_char]
            if plain_segment:
                pdf_text += plain_segment
            ent_segment = text[ent.start_char:ent.end_char]
            if ent_segment:
                color = ENTITY_COLORS.get(ent.label_, "#FFFF00")
                pdf_text += f'<span backColor="{color}">{ent_segment} [{ent.label_}]</span>'
            last_idx = ent.end_char
        if last_idx < len(text):
            pdf_text += text[last_idx:]
        # Generate PDF using ReportLab
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        doc_pdf = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        custom_style = ParagraphStyle(
            'Custom', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=12
        )
        story = []
        story.append(Paragraph(pdf_text, custom_style))
        story.append(Spacer(1, 12))
        doc_pdf.build(story)

    def show_statistics(self):
        """
        Compute and display document statistics:
          1) Word count
          2) Sentence count & avg sentence length
          3) Named entity count & type distribution
          4) Entity density
          5) Named entity length distribution
        Also generates a pie chart and a bar chart using matplotlib.
        """
        raw_text = self.text_input.toPlainText().strip()
        if not raw_text:
            return
        if not any(ch in raw_text for ch in string.punctuation):
            raw_text = preprocess_text(raw_text)
        doc_stat = self.nlp(raw_text)
        total_words = sum(1 for t in doc_stat if not t.is_space and not t.is_punct)
        sentences = list(doc_stat.sents)
        total_sents = len(sentences)
        avg_sent_len = 0.0
        if total_sents > 0:
            words_per_sent = [sum(1 for t in s if not t.is_space and not t.is_punct) for s in sentences]
            avg_sent_len = sum(words_per_sent) / total_sents
        total_ents = len(doc_stat.ents)
        ent_type_counts = Counter([ent.label_ for ent in doc_stat.ents])
        density = (total_ents / total_words * 100) if total_words else 0.0
        ent_len_counts = Counter([len(ent) for ent in doc_stat.ents])
        avg_ent_len = sum(len(ent) for ent in doc_stat.ents) / total_ents if total_ents else 0.0
        lines = [
            f"Total Word Count: {total_words}",
            f"Total Sentence Count: {total_sents}",
            f"Average Sentence Length: {avg_sent_len:.2f} words",
            f"Total Named Entities: {total_ents}",
            f"Named Entity Density: {density:.2f}%",
        ]
        if total_ents > 0:
            lines.append(f"Average Entity Length: {avg_ent_len:.2f} tokens")
            dist_str = ", ".join(f"{lbl}: {cnt}" for lbl, cnt in ent_type_counts.items())
            lines.append(f"Entities by Type: {dist_str}")
            length_dist_str = ", ".join(f"{length} tokens: {count}" for length, count in ent_len_counts.items())
            lines.append(f"Entity Length Distribution: {length_dist_str}")
        else:
            lines.append("Average Entity Length: 0")
            lines.append("Entities by Type: (none)")
            lines.append("Entity Length Distribution: (none)")
        stats_text = "\n".join(lines)
        for i in reversed(range(self.stats_layout.count())):
            w = self.stats_layout.takeAt(i)
            if w.widget():
                w.widget().deleteLater()
        stats_text_widget = QTextEdit(self)
        stats_text_widget.setReadOnly(True)
        stats_text_widget.setPlainText(stats_text)
        self.stats_layout.addWidget(stats_text_widget)
        if total_ents == 0:
            return
        fig1 = Figure(figsize=(5, 4))
        ax1 = fig1.add_subplot(111)
        labels_pie = list(ent_type_counts.keys())
        sizes = [ent_type_counts[k] for k in labels_pie]
        ax1.pie(sizes, labels=labels_pie, autopct="%1.1f%%", startangle=140)
        ax1.set_title("Entity Type Distribution")
        ax1.axis("equal")
        canvas1 = FigureCanvas(fig1)
        self.stats_layout.addWidget(canvas1)
        fig2 = Figure(figsize=(5, 4))
        ax2 = fig2.add_subplot(111)
        sorted_lengths = sorted(ent_len_counts.keys())
        length_counts = [ent_len_counts[l] for l in sorted_lengths]
        ax2.bar(sorted_lengths, length_counts, color="#72B2E4")
        ax2.set_xlabel("Entity Length (tokens)")
        ax2.set_ylabel("Number of Entities")
        ax2.set_title("Named Entity Length Distribution")
        ax2.set_xticks(sorted_lengths)
        canvas2 = FigureCanvas(fig2)
        self.stats_layout.addWidget(canvas2)
        canvas1.draw()
        canvas2.draw()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = NERApp()
    window.show()
    sys.exit(app.exec_())
