from flask import Flask, render_template, request, jsonify
import spacy
import re
import PyPDF2
from docx import Document as DocxDocument

app = Flask(__name__, template_folder="templates")


# 1) Load spaCy model once
nlp = spacy.load("en_core_web_sm")

# 2) Entity labels & colors
LABELS = ["PERSON","ORG","GPE","DATE","TIME","PRODUCT","MONEY",
          "QUANTITY","PHONE","EMAIL","URL","PERCENT"]
COLORS = {
    "PERSON":"#C8E6C9","ORG":"#A2C8EC","GPE":"#FFCDD2","DATE":"#FFF9C4",
    "TIME":"#F8BBD0","PRODUCT":"#B2EBF2","MONEY":"#D1C4E9","QUANTITY":"#C5CAE9",
    "PHONE":"#E0F7FA","EMAIL":"#E1BEE7","URL":"#DCEDC8","PERCENT":"#FFCCBC"
}

def preprocess(text):
    """
    Insert periods if none exist to help spaCy with sentence segmentation.
    """
    if not any(p in text for p in ".!?"):
        tokens = text.split()
        new = []
        for i, t in enumerate(tokens):
            new.append(t)
            if (i < len(tokens) - 1 and
                re.match(r"^\$?\d+(\.\d+)?$", t) and
                tokens[i+1][0].isupper()):
                new.append(".")
        text = " ".join(new)
    return text

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html", labels=LABELS, colors=COLORS)

@app.route("/extract-text", methods=["POST"])
def extract_text():
    """
    Receive a PDF or Word file, extract its text, and return JSON {"text": "..."}.
    """
    uploaded = request.files.get("file")
    if not uploaded or not uploaded.filename:
        return jsonify({"text": ""})

    fname = uploaded.filename.lower()
    text = ""
    if fname.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif fname.endswith((".docx", ".doc")):
        doc = DocxDocument(uploaded)
        text = "\n".join(p.text for p in doc.paragraphs)

    return jsonify({"text": text})

@app.route("/ner", methods=["POST"])
def ner():
    """
    Process either uploaded file or textarea input, run spaCy NER,
    filter by selected entities, and render the highlighted result.
    """
    # 1) Try uploaded file first
    uploaded = request.files.get("file")
    text = ""
    if uploaded and uploaded.filename:
        fname = uploaded.filename.lower()
        if fname.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif fname.endswith((".docx", ".doc")):
            doc = DocxDocument(uploaded)
            text = "\n".join(p.text for p in doc.paragraphs)

    # 2) Fallback to textarea if no file
    if not text:
        text = request.form.get("text", "")

    text = preprocess(text)
    doc = nlp(text)
    selected = request.form.getlist("entities")

    # Build highlighted HTML
    result = []
    last = 0
    for ent in doc.ents:
        if ent.label_ in selected:
            result.append(text[last:ent.start_char])
            color = COLORS.get(ent.label_, "#FFFF00")
            span = (f'<span class="entity" style="background:{color}">'
                    f'{ent.text} [{ent.label_}]</span>')
            result.append(span)
            last = ent.end_char
    result.append(text[last:])
    highlighted = "".join(result)

    return render_template("result.html",
                           labels=LABELS,
                           colors=COLORS,
                           highlighted=highlighted,
                           request=request)

if __name__ == "__main__":
    app.run(debug=True)
